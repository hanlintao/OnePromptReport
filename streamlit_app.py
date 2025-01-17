import streamlit as st
import requests
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import tempfile
import os
import time

# 导入所需的库
from zhipuai import ZhipuAI
from langchain_openai import ChatOpenAI
from langchain_core.prompts import ChatPromptTemplate
from langchain_core.output_parsers import StrOutputParser

# 使用Bing API进行搜索并获取网址
def get_bing_search_results(query, subscription_key, count=10):
    headers = {"Ocp-Apim-Subscription-Key": subscription_key}
    params = {"q": query, "count": count, "textDecorations": True, "textFormat": "HTML"}
    search_url = "https://api.bing.microsoft.com/v7.0/search"
    response = requests.get(search_url, headers=headers, params=params)
    response.raise_for_status()
    search_results = response.json()
    results = [{
        "title": result["name"],
        "url": result["url"],
        "snippet": result["snippet"]
    } for result in search_results["webPages"]["value"]]
    return results

# 打字机效果函数
def typewriter_effect(text, speed=0.005):
    placeholder = st.empty()
    # 添加自定义样式来调整宽度
    custom_css = """
    <style>
    .typewriter-effect {
        width: 100%;
        white-space: pre-wrap;
        word-wrap: break-word;
    }
    </style>
    """
    st.markdown(custom_css, unsafe_allow_html=True)

    current_text = ""
    for char in text:
        current_text += char
        placeholder.markdown(f'<div class="typewriter-effect"><pre>{current_text}</pre></div>', unsafe_allow_html=True)
        time.sleep(speed)

# 设置段落字体为宋体的函数
def set_font_to_songti(paragraph):
    run = paragraph.add_run()
    run.font.name = '宋体'
    r = run._element
    r.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    return run

# 报告生成函数
def generate_report(query, subscription_key, zhipuai_api_key, jina_api_key, prompt1, prompt2, urls, use_gpt4o=False, openai_api_key=None, openai_base_url=None):
    if use_gpt4o:
        if not openai_api_key:
            st.error("请提供OpenAI API Key")
            return None, None
        llm = ChatOpenAI(model="gpt-4o", api_key=openai_api_key, base_url=openai_base_url)
    else:
        client = ZhipuAI(api_key=zhipuai_api_key)

    combined_content = []
    references = []

    progress_text = st.empty()
    progress_bar = st.progress(0)
    total_steps = len(urls)

    for i, url in enumerate(urls, 1):
        progress_text.text(f"正在爬取网址 {i}/{total_steps}: {url}")
        progress_bar.progress(i / total_steps)

        full_url = f"https://r.jina.ai/{url}"
        headers = {
            "Accept": "application/json",
            "Authorization": f"Bearer {jina_api_key}",
            "X-Return-Format": "text"
        }

        try:
            response = requests.get(full_url, headers=headers)
            if response.status_code == 200:
                response_data = response.json()
                content = response_data['data']['text']

                if use_gpt4o:
                    extracted_content = llm.invoke(prompt1.format(content=content)).content
                    #typewriter_effect(extracted_content)
                else:
                    prompt = prompt1.format(content=content)
                    response = client.chat.completions.create(
                        model="glm-4-0520",
                        messages=[
                            {"role": "user", "content": prompt}
                        ],
                    )
                    extracted_content = response.choices[0].message.content
                    #typewriter_effect(extracted_content)

                combined_content.append(extracted_content)
                references.append({
                    "url": url,
                    "summary": extracted_content
                })

            else:
                st.error(f"请求 {full_url} 失败，状态码: {response.status_code}")
        except Exception as e:
            st.error(f"处理 {url} 时发生错误: {e}")

    combined_text = '\n'.join(combined_content)
    report_prompt = prompt2.format(content=combined_text)

    try:
        if use_gpt4o:
            report_content = llm.invoke(report_prompt).content
            #typewriter_effect(report_content)
        else:
            report_response = client.chat.completions.create(
                model="glm-4-0520",
                messages=[
                    {"role": "user", "content": report_prompt}
                ],
            )
            report_content = report_response.choices[0].message.content
            #typewriter_effect(report_content)
    except KeyError as e:
        st.error(f"生成报告时发生错误: {e}")
        st.error(f"报告提示词: {report_prompt}")
        return None, None

    report_document = Document()
    report_document.add_heading('咨询报告', level=1)
    set_font_to_songti(report_document.paragraphs[-1])

    p = report_document.add_paragraph(report_content)
    set_font_to_songti(p)

    report_document.add_heading('附录', level=1)
    set_font_to_songti(report_document.paragraphs[-1])

    for i, reference in enumerate(references, 1):
        report_document.add_heading(f'参考网址 {i}', level=2)
        set_font_to_songti(report_document.paragraphs[-1])

        p = report_document.add_paragraph(reference["url"], style='BodyText')
        set_font_to_songti(p)

        report_document.add_heading('摘要', level=3)
        set_font_to_songti(report_document.paragraphs[-1])

        p = report_document.add_paragraph(reference["summary"], style='BodyText')
        set_font_to_songti(p)

    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp_file:
        temp_filename = tmp_file.name
        report_document.save(temp_filename)

    return report_content, temp_filename

# Streamlit应用程序
st.title("咨询报告一键生成器")

with st.sidebar:
    st.header("选择模型")
    model_choice = st.radio("请选择模型", options=["ZhipuAI", "GPT-4o"])
    use_gpt4o = model_choice == "GPT-4o"

    st.header("API Keys 配置")
    subscription_key = st.text_input("请输入Bing Search API的Subscription Key：", type="password")
    if use_gpt4o:
        openai_api_key = st.text_input("请输入OpenAI API Key：", type="password")
        openai_base_url = st.text_input("请输入OpenAI Base URL：", type="password")
        zhipuai_api_key = None
    else:
        zhipuai_api_key = st.text_input("请输入Zhipu AI的API Key：", type="password")
        openai_api_key = None
        openai_base_url = None

    jina_api_key = st.text_input("请输入Jina API的Key：", type="password")
    
    st.header("Prompt 配置")
    prompt1 = st.text_area("提取内容的Prompt：", value="请提取这个网页中的核心内容，将其转变为一篇主题明确，结构清晰的新的文章：\n\n{content}")
    prompt2 = st.text_area("生成报告的Prompt：", value="你是一位专业的咨询报告撰写人请基于以下内容撰写一份完整的咨询报告，报告应包括标题、摘要、关键词、引言、核心内容、结论和参考文献：\n\n{content}")

    st.header("自定义网址")
    custom_urls = [st.text_input(f"网址 {i+1}：") for i in range(5)]

query = st.text_input("请输入报告主题：")

if st.button("获取搜索结果"):
    if query and subscription_key:
        search_results = get_bing_search_results(query, subscription_key)
        if search_results:
            st.session_state['search_results'] = search_results

if 'search_results' in st.session_state:
    st.header("搜索结果")
    selected_urls = []
    for result in st.session_state['search_results']:
        if st.checkbox(f"{result['title']} ({result['url']})\n{result['snippet']}", key=result['url']):
            selected_urls.append(result['url'])

    if st.button("开始生成报告"):
        if selected_urls:
            report_content, temp_filename = generate_report(query, subscription_key, zhipuai_api_key, jina_api_key, prompt1, prompt2, selected_urls, use_gpt4o, openai_api_key, openai_base_url)
            if report_content:
                st.header("生成的咨询报告")
                typewriter_effect(report_content)
                with open(temp_filename, "rb") as file:
                    btn = st.download_button(
                        label="下载咨询报告",
                        data=file,
                        file_name="consulting_report.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
        else:
            st.error("请选择至少一个搜索结果")
